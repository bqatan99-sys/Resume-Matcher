"""Unit tests for the DOCX resume exporter."""

from io import BytesIO
from zipfile import ZipFile

from docx import Document

from app.services.docx_export import generate_resume_docx_bytes


def test_generate_resume_docx_bytes_contains_expected_sections(sample_resume):
    content = generate_resume_docx_bytes(sample_resume)

    assert content[:2] == b"PK"

    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    assert "JANE DOE" in paragraphs
    assert "PROFESSIONAL SUMMARY" in paragraphs
    assert "EDUCATION" in paragraphs
    assert "EXPERIENCE" in paragraphs
    assert "PROJECTS" in paragraphs
    assert "SKILLS" in paragraphs


def test_generate_resume_docx_bytes_writes_expected_docx_xml(sample_resume):
    content = generate_resume_docx_bytes(sample_resume)

    with ZipFile(BytesIO(content)) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8")

    assert "Times New Roman" in styles_xml
    assert "Senior Backend Engineer" in document_xml
    assert "Backend engineer with 6 years of experience" in document_xml


def test_generate_resume_docx_bytes_can_reuse_template_bytes(sample_resume):
    template = generate_resume_docx_bytes(sample_resume)
    content = generate_resume_docx_bytes(sample_resume, template_bytes=template)

    document = Document(BytesIO(content))

    assert document.sections[0].left_margin == Document(BytesIO(template)).sections[0].left_margin
    assert any(paragraph.text == "EXPERIENCE" for paragraph in document.paragraphs)


def test_generate_resume_docx_bytes_emits_hyperlinks(sample_resume):
    sample_resume["personalInfo"]["website"] = "portfolio.example.com"
    sample_resume["personalInfo"]["linkedin"] = "linkedin.com/in/janedoe"
    sample_resume["personalProjects"][0]["website"] = "demo.example.com"

    content = generate_resume_docx_bytes(sample_resume)

    with ZipFile(BytesIO(content)) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        rels_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")

    assert "w:hyperlink" in document_xml
    assert "https://portfolio.example.com" in rels_xml
    assert "https://linkedin.com/in/janedoe" in rels_xml
    assert "https://demo.example.com" in rels_xml


def test_generate_resume_docx_bytes_keeps_projects_compact(sample_resume):
    sample_resume["personalProjects"][0]["role"] = "Lead PM"
    sample_resume["personalProjects"][0]["years"] = "2024 - Present"
    sample_resume["personalProjects"][0]["github"] = "github.com/janedoe/demo"

    content = generate_resume_docx_bytes(sample_resume)

    document = Document(BytesIO(content))
    project_lines = [paragraph.text for paragraph in document.paragraphs if "API Gateway" in paragraph.text]

    assert project_lines
    assert "Lead PM" not in project_lines[0]
    assert "2024 - Present" not in project_lines[0]
    assert "GitHub" in project_lines[0]
