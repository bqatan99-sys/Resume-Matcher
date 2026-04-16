"""DOCX export helpers with optional source-template reuse."""

from copy import deepcopy
from io import BytesIO
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


NAME_FONT_SIZE = Pt(16)
CONTACT_FONT_SIZE = Pt(9)
BODY_FONT_SIZE = Pt(10)
SECTION_FONT_SIZE = Pt(10)
TOOL_SKILL_HINTS = {
    "sql",
    "python",
    "numpy",
    "pandas",
    "data visualization",
    "lovable",
    "bolt.new",
    "automation",
    "api integration",
    "jira",
    "figma",
    "prototyping",
}


def _set_run_font(run: Any, size: Pt, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = size
    run.bold = bold
    run.italic = italic


def _set_bottom_border(paragraph: Any, color: str = "000000", size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)

    for edge in ("top", "left", "right", "between"):
        existing = borders.find(qn(f"w:{edge}"))
        if existing is None:
            existing = OxmlElement(f"w:{edge}")
            existing.set(qn("w:val"), "none")
            existing.set(qn("w:sz"), "0")
            existing.set(qn("w:space"), "0")
            existing.set(qn("w:color"), "000000")
            borders.append(existing)

    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color)


def _prepare_document(
    template_bytes: bytes | None = None,
) -> tuple[DocumentType, Any, Any | None, int]:
    document = Document(BytesIO(template_bytes)) if template_bytes else Document()
    section = document.sections[0]

    if not template_bytes:
        section.top_margin = Inches(0.39)
        section.bottom_margin = Inches(0.39)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        normal = document.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = BODY_FONT_SIZE

    body = document._element.body
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = deepcopy(child)
        body.remove(child)
    content_width = int(section.page_width - section.left_margin - section.right_margin)
    return document, body, sect_pr, content_width


def _normalize_text(value: str) -> str:
    return " ".join(value.strip().split()).upper()


def _find_heading_index(paragraphs: list[Any], heading: str) -> int | None:
    target = _normalize_text(heading)
    for index, paragraph in enumerate(paragraphs):
        if _normalize_text(paragraph.text) == target:
            return index
    return None


def _next_nonempty(paragraphs: list[Any], start: int, stop: int | None = None) -> Any | None:
    end = len(paragraphs) if stop is None else stop
    for index in range(start, end):
        if paragraphs[index].text.strip():
            return paragraphs[index]
    return None


def _find_blank(paragraphs: list[Any], start: int, stop: int | None = None) -> Any | None:
    end = len(paragraphs) if stop is None else stop
    for index in range(start, end):
        if not paragraphs[index].text.strip():
            return paragraphs[index]
    return None


def _pick_run(paragraph: Any, *, bold: bool | None = None, italic: bool | None = None) -> Any:
    runs = list(paragraph.runs)
    if not runs:
        fallback = paragraph.add_run("")
        return fallback
    for run in runs:
        if bold is not None and bool(run.bold) != bold:
            continue
        if italic is not None and bool(run.italic) != italic:
            continue
        return run
    return runs[0]


def _paragraph_text(paragraph: Any) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def _clone_paragraph_shell(paragraph: Any) -> Any:
    cloned = deepcopy(paragraph._p)
    for child in list(cloned):
        if child.tag != qn("w:pPr"):
            cloned.remove(child)
    return cloned


def _clone_run_xml(source_run: Any, text: str | None = None, tab: bool = False) -> Any:
    cloned = deepcopy(source_run._r)
    for child in list(cloned):
        if child.tag != qn("w:rPr"):
            cloned.remove(child)

    if tab:
        cloned.append(OxmlElement("w:tab"))
        return cloned

    text_element = OxmlElement("w:t")
    if text is None:
        text = ""
    if text.strip() != text or "  " in text:
        text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_element.text = text
    cloned.append(text_element)
    return cloned


def _append_paragraph_xml(body: Any, paragraph_xml: Any) -> None:
    body.append(paragraph_xml)


def _append_simple_paragraph(body: Any, paragraph: Any, text: str, *, bold: bool | None = None, italic: bool | None = None) -> None:
    paragraph_xml = _clone_paragraph_shell(paragraph)
    source_run = _pick_run(paragraph, bold=bold, italic=italic)
    paragraph_xml.append(_clone_run_xml(source_run, text=text))
    _append_paragraph_xml(body, paragraph_xml)


def _append_hyperlink_xml(paragraph_xml: Any, paragraph_part: Any, source_run: Any, text: str, url: str) -> None:
    rel_id = paragraph_part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    hyperlink.append(_clone_run_xml(source_run, text=text))
    paragraph_xml.append(hyperlink)


def _normalize_url(value: str, kind: str = "web") -> str:
    text = value.strip()
    if not text:
        return ""
    lower = text.lower()
    if kind == "email":
        return text if lower.startswith("mailto:") else f"mailto:{text}"
    if kind == "phone":
        return text if lower.startswith("tel:") else f"tel:{text}"
    return text if lower.startswith(("http://", "https://")) else f"https://{text}"


def _build_contact_segments(
    personal_info: dict[str, Any],
    link_labels: dict[str, str] | None = None,
) -> list[tuple[str, str | None]]:
    segments: list[tuple[str, str | None]] = []
    location = (personal_info.get("location") or "").strip()
    email = (personal_info.get("email") or "").strip()
    phone = (personal_info.get("phone") or "").strip()
    linkedin = (personal_info.get("linkedin") or "").strip()
    website = (personal_info.get("website") or "").strip()
    github = (personal_info.get("github") or "").strip()

    if location:
        segments.append((location, None))
    if phone:
        segments.append((phone, _normalize_url(phone, "phone")))
    if email:
        segments.append((email, _normalize_url(email, "email")))
    if linkedin:
        segments.append(((link_labels or {}).get("linkedin") or linkedin, _normalize_url(linkedin)))
    if website:
        segments.append(((link_labels or {}).get("website") or website, _normalize_url(website)))
    if github:
        segments.append(((link_labels or {}).get("github") or github, _normalize_url(github)))
    return segments


def _append_row_paragraph(
    body: Any,
    paragraph: Any,
    left_main: str,
    left_suffix: str = "",
    right_text: str = "",
) -> None:
    paragraph_xml = _clone_paragraph_shell(paragraph)
    bold_run = _pick_run(paragraph, bold=True)
    normal_run = _pick_run(paragraph, bold=False)
    paragraph_xml.append(_clone_run_xml(bold_run, text=left_main))
    if left_suffix:
        paragraph_xml.append(_clone_run_xml(normal_run, text=left_suffix))
    if right_text:
        paragraph_xml.append(_clone_run_xml(bold_run, tab=True))
        paragraph_xml.append(_clone_run_xml(bold_run, text=right_text))
    _append_paragraph_xml(body, paragraph_xml)


def _append_blank_paragraph(body: Any, paragraph: Any | None) -> None:
    if paragraph is None:
        return
    _append_paragraph_xml(body, deepcopy(paragraph._p))


def _classify_contact_link(target: str, text: str) -> str | None:
    lower_target = target.lower()
    lower_text = text.lower()
    if "linkedin" in lower_target or "linkedin" in lower_text:
        return "linkedin"
    if "github" in lower_target or "github" in lower_text:
        return "github"
    if lower_target.startswith("mailto:"):
        return "email"
    if lower_target.startswith("tel:"):
        return "phone"
    return "website"


def _extract_contact_link_labels(paragraph: Any) -> dict[str, str]:
    labels: dict[str, str] = {}
    rels = paragraph.part.rels
    for child in paragraph._p:
        if child.tag != qn("w:hyperlink"):
            continue
        rel_id = child.get(qn("r:id"))
        if not rel_id or rel_id not in rels:
            continue
        target = getattr(rels[rel_id], "target_ref", "") or ""
        text = "".join(node.text or "" for node in child.findall(".//w:t", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})).strip()
        if not text:
            continue
        key = _classify_contact_link(target, text)
        if key and key not in labels:
            labels[key] = text
    return labels


def _extract_skill_rows(paragraphs: list[Any], skills_idx: int) -> list[Any]:
    rows: list[Any] = []
    for paragraph in paragraphs[skills_idx + 1 :]:
        text = _paragraph_text(paragraph)
        normalized = _normalize_text(text)
        if normalized in {"LANGUAGES", "AWARDS", "LINKS", "CERTIFICATIONS"}:
            break
        if text:
            rows.append(paragraph)
    return rows


def _split_technical_skills(skills: list[str]) -> tuple[list[str], list[str]]:
    product: list[str] = []
    tools: list[str] = []
    for skill in skills:
        clean = skill.strip()
        if not clean:
            continue
        if clean.lower() in TOOL_SKILL_HINTS:
            tools.append(clean)
        else:
            product.append(clean)
    if not product and tools:
        return tools, []
    if not tools and product:
        return product, []
    return product, tools


def _build_template_skill_lines(
    additional: dict[str, Any], template_rows: list[Any]
) -> list[str]:
    labels = []
    for row in template_rows:
        text = _paragraph_text(row)
        label = text.split(":", 1)[0].strip() if ":" in text else text.strip()
        if label:
            labels.append(label)

    technical_skills = [str(v).strip() for v in (additional.get("technicalSkills") or []) if str(v).strip()]
    languages = [str(v).strip() for v in (additional.get("languages") or []) if str(v).strip()]
    certifications = [str(v).strip() for v in (additional.get("certificationsTraining") or []) if str(v).strip()]
    awards = [str(v).strip() for v in (additional.get("awards") or []) if str(v).strip()]

    if not labels:
        labels = ["Technical Skills", "Languages", "Certifications", "Awards"]

    if {label.lower() for label in labels[:3]} == {"product", "data & tools", "certifications"}:
        product_skills, tool_skills = _split_technical_skills(technical_skills)
        certification_parts = []
        if certifications:
            certification_parts.extend(certifications)
        if languages:
            certification_parts.append(f"Languages: {', '.join(languages)}")
        if awards:
            certification_parts.append(f"Awards: {', '.join(awards)}")
        mapped = [
            ("Product", product_skills),
            ("Data & Tools", tool_skills),
            ("Certifications", certification_parts),
        ]
        return [f"{label}: {' | '.join(values)}" for label, values in mapped if values]

    generic_rows = [
        ("Technical Skills", technical_skills),
        ("Languages", languages),
        ("Certifications", certifications),
        ("Awards", awards),
    ]
    generic_rows = [(label, values) for label, values in generic_rows if values]
    if len(generic_rows) <= len(labels):
        return [f"{label}: {' | '.join(values)}" for label, values in generic_rows]

    head = generic_rows[: max(0, len(labels) - 1)]
    tail = generic_rows[max(0, len(labels) - 1) :]
    merged_tail = " | ".join(f"{label}: {', '.join(values)}" for label, values in tail)
    lines = [f"{label}: {' | '.join(values)}" for label, values in head]
    if merged_tail:
        final_label = labels[min(len(labels) - 1, len(head))]
        lines.append(f"{final_label}: {merged_tail}")
    return lines


def _extract_template_archetypes(template_bytes: bytes) -> dict[str, Any]:
    template = Document(BytesIO(template_bytes))
    paragraphs = list(template.paragraphs)

    summary_idx = _find_heading_index(paragraphs, "PROFESSIONAL SUMMARY")
    education_idx = _find_heading_index(paragraphs, "EDUCATION")
    experience_idx = _find_heading_index(paragraphs, "EXPERIENCE")
    projects_idx = _find_heading_index(paragraphs, "PROJECTS")
    skills_idx = _find_heading_index(paragraphs, "SKILLS")

    if None in (summary_idx, education_idx, experience_idx, projects_idx, skills_idx):
        raise ValueError("Template headings not found")

    name_para = _next_nonempty(paragraphs, 0, summary_idx)
    contact_para = _next_nonempty(paragraphs, paragraphs.index(name_para) + 1 if name_para else 0, summary_idx)
    summary_body = _next_nonempty(paragraphs, summary_idx + 1, education_idx)
    education_row = _next_nonempty(paragraphs, education_idx + 1, experience_idx)
    education_degree = _next_nonempty(
        paragraphs,
        paragraphs.index(education_row) + 1 if education_row else education_idx + 1,
        experience_idx,
    )
    experience_row = _next_nonempty(paragraphs, experience_idx + 1, projects_idx)
    experience_role = _next_nonempty(
        paragraphs,
        paragraphs.index(experience_row) + 1 if experience_row else experience_idx + 1,
        projects_idx,
    )
    experience_bullet = _next_nonempty(
        paragraphs,
        paragraphs.index(experience_role) + 1 if experience_role else experience_idx + 1,
        projects_idx,
    )
    project_row = _next_nonempty(paragraphs, projects_idx + 1, skills_idx)
    skills_rows = _extract_skill_rows(paragraphs, skills_idx)

    return {
        "name": name_para,
        "contact": contact_para,
        "contact_link_labels": _extract_contact_link_labels(contact_para) if contact_para else {},
        "blank": _find_blank(paragraphs, 0) or _find_blank(paragraphs, summary_idx + 1),
        "summary_heading": paragraphs[summary_idx],
        "summary_body": summary_body,
        "education_heading": paragraphs[education_idx],
        "education_row": education_row,
        "education_degree": education_degree,
        "experience_heading": paragraphs[experience_idx],
        "experience_row": experience_row,
        "experience_role": experience_role,
        "experience_bullet": experience_bullet,
        "projects_heading": paragraphs[projects_idx],
        "project_row": project_row,
        "skills_heading": paragraphs[skills_idx],
        "skills_rows": skills_rows,
    }


def _build_contact_line(personal_info: dict[str, Any]) -> str:
    parts = [text for text, _url in _build_contact_segments(personal_info)]
    return " \u2022 ".join(parts)


def _coerce_lines(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str):
        text = value.strip()
        return [text] if text else []
    return [str(value).strip()]


def _join_inline_text(parts: list[str]) -> str:
    return " ".join(part.strip() for part in parts if part.strip()).strip()


def _project_link_entries(project: dict[str, Any]) -> list[tuple[str, str]]:
    entries: list[tuple[str, str]] = []
    github = str(project.get("github") or "").strip()
    website = str(project.get("website") or "").strip()
    if github:
        entries.append(("GitHub", _normalize_url(github)))
    if website:
        entries.append(("Website", _normalize_url(website)))
    return entries


def _generate_resume_docx_from_template(resume_data: dict[str, Any], template_bytes: bytes) -> bytes:
    document, body, sect_pr, _content_width = _prepare_document(template_bytes)
    archetypes = _extract_template_archetypes(template_bytes)
    personal_info = resume_data.get("personalInfo") or {}

    _append_simple_paragraph(body, archetypes["name"], (personal_info.get("name") or "").strip().upper(), bold=True)
    contact_segments = _build_contact_segments(
        personal_info,
        link_labels=archetypes.get("contact_link_labels") or None,
    )
    if contact_segments:
        paragraph_xml = _clone_paragraph_shell(archetypes["contact"])
        link_run = _pick_run(archetypes["contact"])
        first = True
        for text, url in contact_segments:
            if not first:
                paragraph_xml.append(_clone_run_xml(link_run, text=" \u2022 "))
            first = False
            if url:
                _append_hyperlink_xml(paragraph_xml, archetypes["contact"].part, link_run, text, url)
            else:
                paragraph_xml.append(_clone_run_xml(link_run, text=text))
        _append_paragraph_xml(body, paragraph_xml)
    _append_blank_paragraph(body, archetypes["blank"])

    _append_simple_paragraph(body, archetypes["summary_heading"], "PROFESSIONAL SUMMARY", bold=True)
    summary = str(resume_data.get("summary") or "").strip()
    if summary:
        _append_simple_paragraph(body, archetypes["summary_body"], summary)

    education_items = resume_data.get("education") or []
    if education_items:
        _append_simple_paragraph(body, archetypes["education_heading"], "EDUCATION", bold=True)
        _append_blank_paragraph(body, archetypes["blank"])
        for education in education_items:
            institution = str(education.get("institution") or "").strip()
            location = str(education.get("location") or "").strip()
            years = str(education.get("years") or "").strip()
            left_suffix = f" - {location}" if location else ""
            _append_row_paragraph(body, archetypes["education_row"], institution, left_suffix, years)

            degree = str(education.get("degree") or "").strip()
            extras = _coerce_lines(education.get("description"))
            degree_line = degree
            if extras:
                degree_line = f"{degree_line} - {' | '.join(extras)}" if degree_line else " | ".join(extras)
            if degree_line:
                _append_simple_paragraph(body, archetypes["education_degree"], degree_line, italic=True)
            _append_blank_paragraph(body, archetypes["blank"])

    experience_items = resume_data.get("workExperience") or []
    if experience_items:
        _append_simple_paragraph(body, archetypes["experience_heading"], "EXPERIENCE", bold=True)
        _append_blank_paragraph(body, archetypes["blank"])
        for experience in experience_items:
            company = str(experience.get("company") or "").strip()
            location = str(experience.get("location") or "").strip()
            years = str(experience.get("years") or "").strip()
            left_suffix = f" - {location}" if location else ""
            _append_row_paragraph(body, archetypes["experience_row"], company, left_suffix, years)
            title = str(experience.get("title") or "").strip()
            if title:
                _append_simple_paragraph(body, archetypes["experience_role"], title)
            for line in _coerce_lines(experience.get("description")):
                _append_simple_paragraph(body, archetypes["experience_bullet"], line)
            _append_blank_paragraph(body, archetypes["blank"])

    projects = resume_data.get("personalProjects") or []
    if projects:
        _append_simple_paragraph(body, archetypes["projects_heading"], "PROJECTS", bold=True)
        _append_blank_paragraph(body, archetypes["blank"])
        for project in projects:
            name = str(project.get("name") or "").strip()
            detail_text = _join_inline_text(_coerce_lines(project.get("description")))
            project_links = _project_link_entries(project)
            if name or detail_text or project_links:
                paragraph_xml = _clone_paragraph_shell(archetypes["project_row"])
                bold_run = _pick_run(archetypes["project_row"], bold=True)
                normal_run = _pick_run(archetypes["project_row"], bold=False)
                if name:
                    paragraph_xml.append(_clone_run_xml(bold_run, text=name))
                if detail_text:
                    prefix = ": " if name else ""
                    paragraph_xml.append(_clone_run_xml(normal_run, text=f"{prefix}{detail_text}"))
                for idx, (text, url) in enumerate(project_links):
                    separator = " | " if name or detail_text or idx > 0 else ""
                    if separator:
                        paragraph_xml.append(_clone_run_xml(normal_run, text=separator))
                    _append_hyperlink_xml(paragraph_xml, archetypes["project_row"].part, normal_run, text, url)
                _append_paragraph_xml(body, paragraph_xml)
        _append_blank_paragraph(body, archetypes["blank"])

    additional = resume_data.get("additional") or {}
    skill_lines = _build_template_skill_lines(additional, archetypes.get("skills_rows") or [])
    if skill_lines:
        _append_simple_paragraph(body, archetypes["skills_heading"], _paragraph_text(archetypes["skills_heading"]), bold=True)
        _append_blank_paragraph(body, archetypes["blank"])
        skill_row = (archetypes.get("skills_rows") or [None])[0]
        if skill_row is not None:
            for line in skill_lines:
                _append_simple_paragraph(body, skill_row, line)

    if sect_pr is not None:
        body.append(sect_pr)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_name_header(document: DocumentType, personal_info: dict[str, Any]) -> None:
    name = (personal_info.get("name") or "").strip()
    if not name:
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    _set_bottom_border(paragraph, color="FFFFFF", size="6")

    run = paragraph.add_run(name.upper())
    _set_run_font(run, NAME_FONT_SIZE, bold=True)


def _add_contact_header(document: DocumentType, personal_info: dict[str, Any]) -> None:
    segments = _build_contact_segments(personal_info)
    if not segments:
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    _set_bottom_border(paragraph, color="FFFFFF", size="6")

    first = True
    for text, url in segments:
        if not first:
            sep = paragraph.add_run(" \u2022 ")
            _set_run_font(sep, CONTACT_FONT_SIZE)
        first = False
        if url:
            rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), rel_id)
            run = paragraph.add_run(text)
            _set_run_font(run, CONTACT_FONT_SIZE)
            hyperlink.append(run._r)
            paragraph._p.append(hyperlink)
        else:
            run = paragraph.add_run(text)
            _set_run_font(run, CONTACT_FONT_SIZE)


def _add_spacer(document: DocumentType, size: int = 2) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(size)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.add_run("")


def _add_section_heading(document: DocumentType, heading: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    _set_bottom_border(paragraph)

    run = paragraph.add_run(heading.upper())
    _set_run_font(run, SECTION_FONT_SIZE, bold=True)


def _add_body_paragraph(document: DocumentType, text: str, *, italic: bool = False) -> None:
    if not text.strip():
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text.strip())
    _set_run_font(run, BODY_FONT_SIZE, italic=italic)


def _add_titled_date_row(
    document: DocumentType,
    content_width: int,
    left_main: str,
    left_suffix: str = "",
    right_text: str = "",
) -> None:
    if not left_main and not left_suffix and not right_text:
        return

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)

    if left_main:
        run = paragraph.add_run(left_main)
        _set_run_font(run, BODY_FONT_SIZE, bold=True)
    if left_suffix:
        run = paragraph.add_run(left_suffix)
        _set_run_font(run, BODY_FONT_SIZE)
    if right_text:
        paragraph.add_run("\t")
        run = paragraph.add_run(right_text)
        _set_run_font(run, BODY_FONT_SIZE, bold=True)


def _add_summary_section(document: DocumentType, summary: str) -> None:
    summary = summary.strip()
    if not summary:
        return
    _add_section_heading(document, "Professional Summary")
    _add_spacer(document)
    _add_body_paragraph(document, summary)
    _add_spacer(document)


def _add_education_section(
    document: DocumentType, content_width: int, education_items: list[dict[str, Any]]
) -> None:
    if not education_items:
        return

    _add_section_heading(document, "Education")
    _add_spacer(document)
    for education in education_items:
        institution = (education.get("institution") or "").strip()
        location = (education.get("location") or "").strip()
        years = (education.get("years") or "").strip()
        degree = (education.get("degree") or "").strip()
        left_suffix = f" - {location}" if location else ""
        _add_titled_date_row(document, content_width, institution, left_suffix, years)
        _add_body_paragraph(document, degree, italic=True)
        description = education.get("description")
        if isinstance(description, str) and description.strip():
            _add_body_paragraph(document, description)
        elif isinstance(description, list):
            for line in description:
                _add_body_paragraph(document, str(line))
        _add_spacer(document)


def _add_experience_section(
    document: DocumentType, content_width: int, experiences: list[dict[str, Any]]
) -> None:
    if not experiences:
        return

    _add_section_heading(document, "Experience")
    _add_spacer(document)
    for experience in experiences:
        company = (experience.get("company") or "").strip()
        location = (experience.get("location") or "").strip()
        years = (experience.get("years") or "").strip()
        title = (experience.get("title") or "").strip()
        left_suffix = f" - {location}" if location else ""
        _add_titled_date_row(document, content_width, company, left_suffix, years)
        _add_body_paragraph(document, title)
        for line in experience.get("description") or []:
            _add_body_paragraph(document, str(line))
        _add_spacer(document)


def _add_projects_section(document: DocumentType, projects: list[dict[str, Any]]) -> None:
    if not projects:
        return

    _add_section_heading(document, "Projects")
    _add_spacer(document)
    for project in projects:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1

        name = str(project.get("name") or "").strip()
        if name:
            run = paragraph.add_run(name)
            _set_run_font(run, BODY_FONT_SIZE, bold=True)

        detail_text = _join_inline_text(
            [str(item).strip() for item in (project.get("description") or []) if str(item).strip()]
        )
        if detail_text:
            if name:
                run = paragraph.add_run(": ")
                _set_run_font(run, BODY_FONT_SIZE)
            run = paragraph.add_run(detail_text)
            _set_run_font(run, BODY_FONT_SIZE)
        link_entries = _project_link_entries(project)
        for index, (text, url) in enumerate(link_entries):
            separator = " | " if (name or detail_text or index > 0) else ""
            if separator:
                run = paragraph.add_run(separator)
                _set_run_font(run, BODY_FONT_SIZE)
            rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), rel_id)
            run = paragraph.add_run(text)
            _set_run_font(run, BODY_FONT_SIZE)
            hyperlink.append(run._r)
            paragraph._p.append(hyperlink)
    _add_spacer(document)


def _add_additional_section(document: DocumentType, additional: dict[str, Any]) -> None:
    if not additional:
        return

    rows = [
        ("Technical Skills", additional.get("technicalSkills") or []),
        ("Languages", additional.get("languages") or []),
        ("Certifications", additional.get("certificationsTraining") or []),
        ("Awards", additional.get("awards") or []),
    ]
    rows = [(label, values) for label, values in rows if values]
    if not rows:
        return

    _add_section_heading(document, "Skills")
    _add_spacer(document)
    for label, values in rows:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        run = paragraph.add_run(f"{label}: ")
        _set_run_font(run, BODY_FONT_SIZE, bold=True)
        run = paragraph.add_run(
            " | ".join(str(value).strip() for value in values if str(value).strip())
        )
        _set_run_font(run, BODY_FONT_SIZE)


def _add_custom_sections(
    document: DocumentType, content_width: int, resume_data: dict[str, Any]
) -> None:
    custom_sections = resume_data.get("customSections") or {}
    section_meta = resume_data.get("sectionMeta") or []
    for meta in sorted(section_meta, key=lambda item: item.get("order", 999)):
        if meta.get("isDefault") or not meta.get("isVisible"):
            continue
        key = meta.get("key")
        section = custom_sections.get(key)
        if not section:
            continue

        _add_section_heading(document, meta.get("displayName") or key or "Section")
        _add_spacer(document)
        section_type = section.get("sectionType")
        if section_type == "text":
            _add_body_paragraph(document, str(section.get("text") or ""))
        elif section_type == "stringList":
            for item in section.get("strings") or []:
                _add_body_paragraph(document, str(item))
        elif section_type == "itemList":
            for item in section.get("items") or []:
                title = str(item.get("title") or "").strip()
                subtitle = str(item.get("subtitle") or "").strip()
                location = str(item.get("location") or "").strip()
                years = str(item.get("years") or "").strip()
                left_suffix = f" - {location}" if location else ""
                _add_titled_date_row(document, content_width, title, left_suffix, years)
                if subtitle:
                    _add_body_paragraph(document, subtitle, italic=True)
                for line in item.get("description") or []:
                    _add_body_paragraph(document, str(line))
                _add_spacer(document)
        _add_spacer(document)


def generate_resume_docx_bytes(
    resume_data: dict[str, Any], template_bytes: bytes | None = None
) -> bytes:
    """Build a DOCX resume, reusing the original uploaded template when available."""
    if template_bytes:
        try:
            return _generate_resume_docx_from_template(resume_data, template_bytes)
        except Exception:
            pass

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.39)
    section.bottom_margin = Inches(0.39)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = BODY_FONT_SIZE
    content_width = int(section.page_width - section.left_margin - section.right_margin)

    personal_info = resume_data.get("personalInfo") or {}
    _add_name_header(document, personal_info)
    _add_contact_header(document, personal_info)
    _add_spacer(document, size=6)

    _add_summary_section(document, str(resume_data.get("summary") or ""))
    _add_education_section(document, content_width, resume_data.get("education") or [])
    _add_experience_section(document, content_width, resume_data.get("workExperience") or [])
    _add_projects_section(document, resume_data.get("personalProjects") or [])
    _add_additional_section(document, resume_data.get("additional") or {})
    _add_custom_sections(document, content_width, resume_data)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
